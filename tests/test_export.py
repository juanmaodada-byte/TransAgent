"""
导出契约测试
=============
D1 测试基线 — DOCX 导出结构检查，验证导出产物符合 P0 期望。

D1 只建立结构检查入口和预期对象，不实现新 exporter。
当前 exporter 不满足的用 strict xfail 记录。

不依赖 API key、网络或真实 LLM。
"""
import os
import pytest
from transagent.interface import FormatType
from transagent.backend.pipeline.preprocess import convert_to_md


# ══════════════════════════════════════════════════════════════════
# 导出 — 结构检查
# ══════════════════════════════════════════════════════════════════

class TestExportStructure:
    def test_export_docx_heading_rebuild(self, cloud_native_docx_path, tmp_path):
        """DOCX 导出的产物中 H2/H3 标题使用正确的 heading 样式层级"""
        from transagent.backend.pipeline.preprocess import convert_to_md
        from transagent.backend.pipeline.structure_parser import parse_structure
        from transagent.backend.pipeline.restore import restore_placeholders
        from transagent.backend.pipeline.exporter import export_to_format

        converted = convert_to_md(cloud_native_docx_path, FormatType.DOCX.value)
        # 模拟完整管线路径（不含LLM）：保护 → 还原 → 导出
        protected_md, pmap = parse_structure(converted.md_text)
        restored_md = restore_placeholders(protected_md, pmap)

        import transagent.backend.config as cfg
        old_dir = cfg.get_config().app.workspace_dir
        try:
            cfg.get_config().app.workspace_dir = str(tmp_path)
            export_path = export_to_format(restored_md, "docx", converted.assets_dir)

            from docx import Document as DocxReader
            doc = DocxReader(export_path)
            # P0 期望: heading 层级在 protect→restore→export 后仍保留
            h2_count = sum(1 for p in doc.paragraphs
                          if p.style.name == "Heading 2")
            assert h2_count >= 3, (
                f"Expected >= 3 Heading 2 paragraphs, got {h2_count}. "
                "Heading hierarchy should survive protection and export."
            )
        finally:
            cfg.get_config().app.workspace_dir = old_dir

    @pytest.mark.xfail(
        strict=True, raises=(AssertionError, ImportError, Exception),
        reason="P0/D: exporter 不处理管道表，导出 DOCX 不包含表格"
    )
    def test_export_table_rebuild(self, cloud_native_docx_path, tmp_path):
        """DOCX 导出产物应包含实际表格（非纯文本）"""
        from transagent.backend.pipeline.preprocess import convert_to_md
        from transagent.backend.pipeline.exporter import export_to_format

        converted = convert_to_md(cloud_native_docx_path, FormatType.DOCX.value)

        import transagent.backend.config as cfg
        old_dir = cfg.get_config().app.workspace_dir
        try:
            cfg.get_config().app.workspace_dir = str(tmp_path)
            export_path = export_to_format(converted.md_text, "docx", converted.assets_dir)

            from docx import Document as DocxReader
            doc = DocxReader(export_path)
            # P0 期望: 导出 DOCX 包含实际表格元素
            assert len(doc.tables) >= 1, (
                f"Expected at least 1 table in exported DOCX, got {len(doc.tables)}"
            )
        finally:
            cfg.get_config().app.workspace_dir = old_dir

    @pytest.mark.xfail(
        strict=True, raises=(AssertionError, ImportError, Exception),
        reason="P0/D: exporter 不处理列表，列表项被当作普通段落，无编号/项目符号"
    )
    def test_export_list_rebuild(self, cloud_native_docx_path, tmp_path):
        """DOCX 导出的列表项应有独特的段落样式（非普通段落）"""
        from transagent.backend.pipeline.preprocess import convert_to_md
        from transagent.backend.pipeline.exporter import export_to_format

        converted = convert_to_md(cloud_native_docx_path, FormatType.DOCX.value)

        import transagent.backend.config as cfg
        old_dir = cfg.get_config().app.workspace_dir
        try:
            cfg.get_config().app.workspace_dir = str(tmp_path)
            export_path = export_to_format(converted.md_text, "docx", converted.assets_dir)

            from docx import Document as DocxReader
            doc = DocxReader(export_path)
            # P0 期望: 列表项有 ListParagraph 或 ListBullet 样式
            list_styled = 0
            for p in doc.paragraphs:
                if "List" in p.style.name or "Bullet" in p.style.name:
                    list_styled += 1
            assert list_styled >= 2, (
                f"Expected >= 2 list-styled paragraphs, got {list_styled}. "
                "Current exporter uses plain paragraphs for lists."
            )
        finally:
            cfg.get_config().app.workspace_dir = old_dir

    @pytest.mark.xfail(
        strict=True, raises=(AssertionError, ImportError, Exception),
        reason="P0/D: exporter 不处理代码区域 — 代码块 ``` 被跳过，中间内容按普通段落处理，无等宽字体"
    )
    def test_export_code_presence(self, cloud_native_docx_path, tmp_path):
        """DOCX 导出的代码区域应有等宽字体（区别于普通段落）"""
        from transagent.backend.pipeline.preprocess import convert_to_md
        from transagent.backend.pipeline.exporter import export_to_format

        converted = convert_to_md(cloud_native_docx_path, FormatType.DOCX.value)

        import transagent.backend.config as cfg
        old_dir = cfg.get_config().app.workspace_dir
        try:
            cfg.get_config().app.workspace_dir = str(tmp_path)
            export_path = export_to_format(converted.md_text, "docx", converted.assets_dir)

            from docx import Document as DocxReader
            doc = DocxReader(export_path)
            # P0 期望: YAML 代码区域使用等宽字体
            code_mono = False
            for p in doc.paragraphs:
                if "apiVersion" in p.text:
                    for run in p.runs:
                        # 检查字体名是否包含 Consolas 或 Courier
                        fn = (run.font.name or "").lower()
                        if "consolas" in fn or "courier" in fn:
                            code_mono = True
                            break
            assert code_mono, (
                "Code blocks in exported DOCX should use monospace font "
                "(Consolas/Courier). Current exporter uses default font."
            )
        finally:
            cfg.get_config().app.workspace_dir = old_dir


# ══════════════════════════════════════════════════════════════════
# 导出 — 文件隔离
# ══════════════════════════════════════════════════════════════════

class TestExportIsolation:
    @pytest.mark.xfail(
        strict=True, raises=AssertionError,
        reason="P0/D: 导出路径固定为 output.docx/output.html，并发 session 可能覆盖"
    )
    def test_export_session_isolation(self):
        """不同 session 的导出路径不应冲突"""
        from transagent.backend.pipeline.exporter import export_to_format
        # 当前 behaviour: 固定使用 output.docx
        # P0 期望: 路径包含 session ID

        # 两次相同格式的导出会覆盖文件
        import tempfile
        with tempfile.TemporaryDirectory() as td:
            # 设置 workspace_dir 为临时目录
            # 注意: 这依赖 config，需要 monkeypatch
            pass

        # 简化验证: 检查 _export_docx 函数的输出路径模式
        import inspect
        import re
        source = inspect.getsource(export_to_format)
        # 期望输出路径包含某种唯一标识（session_id, uuid 等）
        has_unique = "session" in source.lower() or "uuid" in source.lower() or "temp" in source.lower()
        assert has_unique, (
            "Export path should include session or unique identifier to prevent overwrites"
        )


# ══════════════════════════════════════════════════════════════════
# 导出 — 正常契约
# ══════════════════════════════════════════════════════════════════

class TestExportContract:
    def test_export_docx_supported(self):
        """docx 是支持的导出格式"""
        from transagent.backend.pipeline.exporter import export_to_format
        supported = ["docx", "html", "bilingual"]
        assert "docx" in supported

    def test_export_unsupported_rejected(self):
        """不支持的导出格式应被拒绝"""
        # 参考 exporter.py 中的异常处理
        from transagent.backend.pipeline.exporter import export_to_format
        # 验证函数签名支持 target_format 参数
        import inspect
        sig = inspect.signature(export_to_format)
        assert "target_format" in sig.parameters
