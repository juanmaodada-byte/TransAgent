from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from transagent.backend.pipeline import pdf_readability as readability


SOURCE_TEXT = "Alpha anchor one. Beta anchor two. Gamma anchor three. Delta anchor four. Epsilon anchor five."


def _write_docx(path: Path, images: int = 0, tables: int = 0) -> None:
    doc = Document()
    doc.add_paragraph("normalized")
    for _ in range(tables):
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "table"
    doc.save(path)


def _patch_audit(monkeypatch, tmp_path, *, source_text=SOURCE_TEXT, rendered_text=SOURCE_TEXT, source_pages=1, rendered_pages=1, images=0, tables=0, column=False):
    pdf = tmp_path / "source.pdf"
    pdf.write_bytes(b"%PDF-1.7\n%%EOF")
    docx = tmp_path / "normalized.docx"
    _write_docx(docx)
    monkeypatch.setattr(
        readability,
        "_inspect_source_pdf",
        lambda *args, **kwargs: {
            "page_count": source_pages,
            "pages_text": [source_text],
            "image_count": images,
            "table_count": tables,
            "column_layout_risk": column,
        },
    )
    monkeypatch.setattr(readability, "render_docx_to_pdf", lambda docx_path, out_dir: tmp_path / "rendered.pdf")
    monkeypatch.setattr(readability, "pdf_page_count", lambda pdf_path: rendered_pages)
    png = tmp_path / "page.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\nnot-real-but-mocked")
    monkeypatch.setattr(readability, "render_pdf_to_pngs", lambda pdf_path, out_dir, pages: [png] * rendered_pages)
    monkeypatch.setattr(readability, "is_png_nonblank", lambda path: True)
    monkeypatch.setattr(readability, "extract_pdf_text", lambda pdf_path, out_dir: rendered_text)
    monkeypatch.setattr(
        readability.pdf_overlap,
        "detect_rendered_text_image_overlap",
        lambda pdf_path: {"text_image_overlap_count": 0, "text_image_overlap_max_ratio": 0.0},
    )
    monkeypatch.setattr(
        readability,
        "check_pdf_image_visibility",
        lambda ref_pdf, cand_pdf: {
            "image_visibility_checked": True,
            "meaningful_image_count": 0,
            "matched_visible_image_count": 0,
            "minimum_visible_area_ratio": 1.0,
            "invisible_image_count": 0,
        },
    )
    return pdf, docx


def test_plain_pdf_content_retention_is_normal(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path)
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    assert result.metadata["content_retention_ratio"] >= 0.90
    assert result.metadata["length_coverage_ratio"] == 1.0
    assert result.metadata["fuzzy_token_coverage_ratio"] == 1.0
    assert result.warnings == []


def test_minor_character_noise_is_tolerated(monkeypatch, tmp_path):
    source = "The ordinary document keeps AlphaToken BetaToken GammaToken in sequence."
    rendered = "The ordnary document keeps AlphaToken BetaToken GammaToken in sequence."
    pdf, docx = _patch_audit(monkeypatch, tmp_path, source_text=source, rendered_text=rendered)
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    assert result.metadata["content_retention_ratio"] >= 0.90
    assert "content retention" not in " ".join(result.warnings)


def test_repeated_text_with_all_unique_tokens_still_blocks(monkeypatch, tmp_path):
    unique_tokens = "AlphaToken BetaToken GammaToken DeltaToken EpsilonToken"
    source = " ".join([unique_tokens] * 80)
    rendered = unique_tokens
    assert readability._token_coverage(readability._normalize_text(source), readability._normalize_text(rendered)) == 1.0
    assert readability._length_coverage(readability._normalize_text(source), readability._normalize_text(rendered)) < 0.60
    pdf, docx = _patch_audit(monkeypatch, tmp_path, source_text=source, rendered_text=rendered)
    with pytest.raises(ValueError, match="DOCUMENT_INTEGRITY_ERROR"):
        readability.audit_pdf_readability(pdf, docx, tmp_path)


def test_seventy_percent_content_warns_without_blocking(monkeypatch, tmp_path):
    source = " ".join(f"Token{index:03d}" for index in range(100))
    rendered = " ".join(f"Token{index:03d}" for index in range(70))
    pdf, docx = _patch_audit(monkeypatch, tmp_path, source_text=source, rendered_text=rendered)
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    assert 0.60 <= result.metadata["content_retention_ratio"] < 0.90
    assert any("content retention ratio" in warning for warning in result.warnings)


@pytest.mark.parametrize(
    ("source_text", "rendered_text"),
    [
        ("", ""),
        ("AlphaToken BetaToken", ""),
        ("", "AlphaToken BetaToken"),
    ],
)
def test_empty_source_or_rendered_text_has_zero_retention(monkeypatch, tmp_path, source_text, rendered_text):
    source = readability._normalize_text(source_text)
    rendered = readability._normalize_text(rendered_text)
    assert readability._length_coverage(source, rendered) == 0.0
    pdf, docx = _patch_audit(monkeypatch, tmp_path, source_text=source_text, rendered_text=rendered_text)
    with pytest.raises(ValueError, match="DOCUMENT_INTEGRITY_ERROR"):
        readability.audit_pdf_readability(pdf, docx, tmp_path)


def test_double_column_warning(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path, column=True)
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    assert readability.READING_ORDER_WARNING in result.warnings


def test_page_count_change_records_warning(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path, source_pages=1, rendered_pages=2)
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    assert "changed from 1 to 2" in " ".join(result.warnings)


def test_severe_content_loss_blocks(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path, rendered_text="tiny")
    with pytest.raises(ValueError, match="DOCUMENT_INTEGRITY_ERROR"):
        readability.audit_pdf_readability(pdf, docx, tmp_path)


def test_text_image_overlap_blocks_pdf_delivery(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path)
    monkeypatch.setattr(
        readability.pdf_overlap,
        "detect_rendered_text_image_overlap",
        lambda pdf_path: {"text_image_overlap_count": 1, "text_image_overlap_max_ratio": 0.42},
    )
    with pytest.raises(ValueError, match="DOCUMENT_INTEGRITY_ERROR"):
        readability.audit_pdf_readability(pdf, docx, tmp_path)


def test_reading_order_same_order_is_zero():
    anchors = [f"Anchor {index}" for index in range(1, 7)]
    rendered = " ".join(anchors)
    assert readability._reading_order_inversion_ratio(anchors, rendered) == 0.0


def test_reading_order_reverse_order_is_one():
    anchors = [f"Anchor {index}" for index in range(1, 7)]
    rendered = " ".join(reversed(anchors))
    assert readability._reading_order_inversion_ratio(anchors, rendered) == 1.0


def test_reading_order_rotation_exceeds_warning_threshold():
    anchors = [f"Anchor {index}" for index in range(1, 7)]
    rendered = " ".join([anchors[1], anchors[2], anchors[3], anchors[4], anchors[5], anchors[0]])
    assert readability._reading_order_inversion_ratio(anchors, rendered) > readability.ORDER_WARNING_THRESHOLD


def test_reading_order_ignores_missing_anchors():
    anchors = [f"Anchor {index}" for index in range(1, 5)]
    rendered = "Anchor 3 Anchor 1"
    metrics = readability._reading_order_metrics(anchors, rendered)
    assert metrics["found_anchor_count"] == 2
    assert metrics["comparable_anchor_pairs"] == 1
    assert metrics["reading_order_inversion_ratio"] == 1.0


def test_image_loss_warning(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path, images=2)
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    assert readability.OBJECT_LOSS_WARNING in result.warnings


def test_metadata_does_not_contain_full_text(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path)
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    assert SOURCE_TEXT not in repr(result.metadata)
    assert "content_retention_ratio" in result.metadata
    assert "source_text" not in result.metadata
    assert "rendered_text" not in result.metadata


def test_layer1_image_visibility_metadata_recorded(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path, images=2)
    monkeypatch.setattr(
        readability,
        "check_pdf_image_visibility",
        lambda ref_pdf, cand_pdf: {
            "image_visibility_checked": True,
            "meaningful_image_count": 2,
            "matched_visible_image_count": 2,
            "minimum_visible_area_ratio": 1.0,
            "invisible_image_count": 0,
        },
    )
    result = readability.audit_pdf_readability(pdf, docx, tmp_path)
    image_visibility = result.metadata["image_visibility"]
    assert image_visibility["image_visibility_checked"] is True
    assert image_visibility["meaningful_image_count"] == 2
    assert image_visibility["matched_visible_image_count"] == 2
    assert image_visibility["invisible_image_count"] == 0


def test_layer1_blocks_when_normalization_loses_a_figure(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path, images=1)
    monkeypatch.setattr(
        readability,
        "check_pdf_image_visibility",
        lambda ref_pdf, cand_pdf: {
            "image_visibility_checked": True,
            "meaningful_image_count": 1,
            "matched_visible_image_count": 0,
            "minimum_visible_area_ratio": 0.0,
            "invisible_image_count": 1,
        },
    )
    with pytest.raises(ValueError, match="DOCUMENT_INTEGRITY_ERROR"):
        readability.audit_pdf_readability(pdf, docx, tmp_path)


def test_layer1_runtime_unavailable_maps_to_runtime_error(monkeypatch, tmp_path):
    pdf, docx = _patch_audit(monkeypatch, tmp_path, images=1)
    def boom(ref_pdf, cand_pdf):
        raise readability.quality_error("DOCUMENT_RUNTIME_UNAVAILABLE", "PDF runtime unavailable")

    monkeypatch.setattr(readability, "check_pdf_image_visibility", boom)
    with pytest.raises(ValueError, match="DOCUMENT_RUNTIME_UNAVAILABLE"):
        readability.audit_pdf_readability(pdf, docx, tmp_path)
