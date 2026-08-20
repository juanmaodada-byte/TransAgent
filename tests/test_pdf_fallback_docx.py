"""D6 fallback text DOCX tests."""

from __future__ import annotations

from pathlib import Path
from xml.etree import ElementTree as ET
import zipfile

from docx import Document

from backend.pipeline.pdf_probe import FALLBACK_WARNING, ensure_pdf_fixtures, fallback_text_docx, validate_docx_package


def _document_xml(docx_path: Path) -> ET.Element:
    with zipfile.ZipFile(docx_path) as zf:
        return ET.fromstring(zf.read("word/document.xml"))


def test_fallback_writes_pages_in_order(tmp_path):
    pdf_dir = ensure_pdf_fixtures()
    output = tmp_path / "fallback.docx"
    result = fallback_text_docx(pdf_dir / "d6_plain_text.pdf", output)
    doc = Document(output)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Page 1 repeats" in text
    assert "Page 2 repeats" in text
    assert text.index("Page 1 repeats") < text.index("Page 2 repeats")
    assert result.fallback_used


def test_fallback_inserts_page_breaks(tmp_path):
    pdf_dir = ensure_pdf_fixtures()
    output = tmp_path / "fallback.docx"
    fallback_text_docx(pdf_dir / "d6_plain_text.pdf", output)
    xml = zipfile.ZipFile(output).read("word/document.xml").decode("utf-8")
    assert 'w:type="page"' in xml


def test_fallback_does_not_create_empty_paragraph_explosion(tmp_path):
    pdf_dir = ensure_pdf_fixtures()
    output = tmp_path / "fallback.docx"
    fallback_text_docx(pdf_dir / "d6_plain_text.pdf", output)
    doc = Document(output)
    empty = sum(1 for paragraph in doc.paragraphs if not paragraph.text)
    assert empty <= 2
    assert len(doc.paragraphs) < 40


def test_fallback_warning_mentions_layout_and_non_text_loss(tmp_path):
    pdf_dir = ensure_pdf_fixtures()
    output = tmp_path / "fallback.docx"
    result = fallback_text_docx(pdf_dir / "d6_plain_text.pdf", output)
    assert result.warnings == [FALLBACK_WARNING]
    assert "layout" in result.warnings[0]
    assert "non-text content loss" in result.warnings[0]


def test_fallback_docx_is_parseable(tmp_path):
    pdf_dir = ensure_pdf_fixtures()
    output = tmp_path / "fallback.docx"
    fallback_text_docx(pdf_dir / "d6_plain_text.pdf", output)
    package = validate_docx_package(output)
    assert package["xml_file_count"] > 0
    assert _document_xml(output).tag.endswith("document")
