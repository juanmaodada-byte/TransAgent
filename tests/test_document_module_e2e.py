from __future__ import annotations

from pathlib import Path
from xml.etree import ElementTree as ET
import zipfile

import pytest
from docx import Document

from transagent.interface import DocumentBlock
from transagent.backend.pipeline import native_document as native
from transagent.backend.pipeline import pdf_normalizer
from transagent.backend.pipeline.doc_normalizer import resolve_libreoffice
from transagent.backend.pipeline.docx_snapshot import snapshot_docx_structure
from transagent.backend.pipeline.document_quality import check_document_runtime_health
from backend.pipeline.pdf_probe import ensure_pdf_fixtures


def _runtime_ready() -> bool:
    health = check_document_runtime_health()
    return bool(health["tikal_1_48_0"]["ok"] and health["okapi_config"]["ok"] and health["libreoffice"]["ok"] and health["pdfinfo"]["ok"] and health["pdftoppm"]["ok"])


def _translated(blocks: list[DocumentBlock]) -> list[DocumentBlock]:
    return [DocumentBlock(block_id=block.block_id, block_type=block.block_type, text=f"D8译:{block.text}") for block in blocks]


def _merge_or_skip_cjk_env(doc, blocks: list[DocumentBlock], session_id: str) -> Path:
    try:
        return Path(native.merge_translations(doc, blocks, session_id=session_id))
    except ValueError as exc:
        if str(exc).startswith("DOCUMENT_RUNTIME_UNAVAILABLE: CJK font rendering runtime is unavailable"):
            pytest.skip("CJK font rendering runtime unavailable")
        raise


def _assert_successful_merge(doc, output: Path) -> None:
    assert doc.blocks
    assert len({block.block_id for block in doc.blocks}) == len(doc.blocks)
    assert output.exists() and output.suffix == ".docx"
    assert output.resolve() != Path(doc.source_document_path).resolve()
    assert output.resolve() != Path(doc.normalized_docx_path).resolve()
    assert snapshot_docx_structure(str(output)) == doc.original_structure_snapshot
    assert doc.document_manifest is not None
    assert doc.document_manifest.delivery_quality["page_count"] > 0
    assert doc.document_manifest.delivery_quality["blank_pages"] == []
    assert doc.document_manifest.delivery_quality["u_fffd_found"] is False
    with zipfile.ZipFile(output) as zf:
        text = "".join(ET.fromstring(zf.read("word/document.xml")).itertext())
        assert "\ufffd" not in text
        assert "[[TA_" not in text


def test_fallback_docx_readable_style_contract(tmp_path):
    pdf_dir = ensure_pdf_fixtures()
    output = tmp_path / "fallback.docx"
    pdf_normalizer._fallback_text_docx(pdf_dir / "d6_plain_text.pdf", output)
    doc = Document(output)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "https://example.com/transagent/d6/pdf-probe?case=plain" in text
    assert "python3 -m pytest tests/test_pdf_probe.py -q" in text
    empty = sum(1 for paragraph in doc.paragraphs if not paragraph.text)
    assert empty <= 2
    section = doc.sections[0]
    for margin in [section.top_margin, section.bottom_margin, section.left_margin, section.right_margin]:
        assert 0.70 <= margin.inches <= 1.05
    normal = doc.styles["Normal"]
    assert 10 <= normal.font.size.pt <= 11
    assert 1.10 <= normal.paragraph_format.line_spacing <= 1.20
    assert 4 <= normal.paragraph_format.space_after.pt <= 6
    xml = zipfile.ZipFile(output).read("word/document.xml").decode("utf-8")
    assert 'w:type="page"' in xml


@pytest.mark.integration
def test_docx_native_e2e(okapi_probe_docx_path):
    if not _runtime_ready():
        pytest.skip("document runtime unavailable")
    doc = native.extract_document(okapi_probe_docx_path, session_id="d8docx")
    assert doc.fidelity_level == "native"
    output = _merge_or_skip_cjk_env(doc, _translated(doc.blocks), "d8docx")
    _assert_successful_merge(doc, output)


@pytest.mark.integration
def test_doc_normalized_e2e(real_word_97_doc_path):
    if not _runtime_ready():
        pytest.skip("document runtime unavailable")
    resolve_libreoffice()
    doc = native.extract_document(real_word_97_doc_path, session_id="d8doc")
    assert doc.fidelity_level == "normalized"
    assert doc.conversion_warnings
    output = _merge_or_skip_cjk_env(doc, _translated(doc.blocks), "d8doc")
    _assert_successful_merge(doc, output)


@pytest.mark.integration
def test_plain_pdf_pdf2docx_e2e():
    if not _runtime_ready():
        pytest.skip("document runtime unavailable")
    pdf_dir = ensure_pdf_fixtures()
    doc = native.extract_document(str(pdf_dir / "d6_plain_text.pdf"), session_id="d8pdfplain")
    assert doc.fidelity_level == "approximate"
    assert doc.document_manifest.conversion_metadata["fallback_used"] is False
    assert "readability" in doc.document_manifest.conversion_metadata
    output = Path(native.merge_translations(doc, _translated(doc.blocks), session_id="d8pdfplain"))
    _assert_successful_merge(doc, output)


@pytest.mark.integration
def test_double_column_pdf_warning_e2e():
    if not _runtime_ready():
        pytest.skip("document runtime unavailable")
    pdf_dir = ensure_pdf_fixtures()
    doc = native.extract_document(str(pdf_dir / "d6_double_column.pdf"), session_id="d8pdfcolumns")
    joined = "\n".join(doc.conversion_warnings)
    assert "PDF column reading order may require manual review." in joined
    assert doc.document_manifest.conversion_metadata["readability"]["reading_order_inversion_ratio"] >= 0


@pytest.mark.integration
def test_mixed_pdf_warning_e2e():
    if not _runtime_ready():
        pytest.skip("document runtime unavailable")
    pdf_dir = ensure_pdf_fixtures()
    doc = native.extract_document(str(pdf_dir / "d6_mixed_text_and_scan.pdf"), session_id="d8pdfmixed")
    assert pdf_normalizer.MIXED_TEXT_WARNING in doc.conversion_warnings
    assert doc.document_manifest.conversion_metadata["no_text_pages"]


@pytest.mark.integration
def test_forced_fallback_pdf_e2e(monkeypatch):
    if not _runtime_ready():
        pytest.skip("document runtime unavailable")
    pdf_dir = ensure_pdf_fixtures()
    monkeypatch.setattr(
        pdf_normalizer,
        "_convert_with_pdf2docx",
        lambda *args: (_ for _ in ()).throw(pdf_normalizer.pdf_normalizer_error("DOCUMENT_CONVERSION_ERROR", "PDF to DOCX conversion failed")),
    )
    monkeypatch.setattr(native, "normalize_pdf_to_docx", pdf_normalizer.normalize_pdf_to_docx)
    doc = native.extract_document(str(pdf_dir / "d6_plain_text.pdf"), session_id="d8fallback")
    assert doc.document_manifest.conversion_metadata["fallback_used"] is True
    assert pdf_normalizer.FALLBACK_WARNING in doc.conversion_warnings
    output = Path(native.merge_translations(doc, _translated(doc.blocks), session_id="d8fallback"))
    _assert_successful_merge(doc, output)


@pytest.mark.integration
def test_scanned_pdf_rejected_e2e():
    pdf_dir = ensure_pdf_fixtures()
    with pytest.raises(ValueError, match="DOCUMENT_OCR_UNSUPPORTED"):
        native.extract_document(str(pdf_dir / "d6_scanned_image_only.pdf"), session_id="d8scan")
