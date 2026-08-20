"""Deterministic fixture coverage tests."""

from __future__ import annotations

import hashlib
import os
import subprocess
import sys

from docx import Document

from tests.conftest import FIXTURES_DIR


FIXTURE_FILES = [
    "plain_technical.docx",
    "api_reference.docx",
    "cloud_native_mixed.docx",
    "okapi_probe_mixed.docx",
    "format_exceptions/pdf_named_docx.docx",
    "format_exceptions/corrupt.docx",
    "format_exceptions/plain_zip.docx",
    "format_exceptions/minimal.pdf",
    "format_exceptions/unsupported.bin",
    "format_exceptions/fake_ole.doc",
    "format_exceptions/minimal_word.doc",
    "format_exceptions/workbook_named_doc.doc",
    "format_exceptions/encrypted_ole.doc",
]


def _sha256(path):
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()


def test_fixture_generator_is_deterministic():
    before = {
        name: _sha256(os.path.join(FIXTURES_DIR, name))
        for name in FIXTURE_FILES
    }
    script = os.path.join(FIXTURES_DIR, "generate_docx_fixture.py")
    subprocess.run([sys.executable, script], check=True, cwd=os.path.dirname(FIXTURES_DIR))
    after = {
        name: _sha256(os.path.join(FIXTURES_DIR, name))
        for name in FIXTURE_FILES
    }
    assert after == before


def test_plain_docx_fixture_covers_runs_and_list(plain_technical_docx_path):
    doc = Document(plain_technical_docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Plain Technical Guide" in text
    assert any(run.bold for p in doc.paragraphs for run in p.runs)
    assert any(run.italic for p in doc.paragraphs for run in p.runs)
    assert sum(1 for p in doc.paragraphs if "List" in p.style.name) >= 3


def test_api_docx_fixture_covers_api_risks(api_reference_docx_path):
    doc = Document(api_reference_docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "https://api.example.com/v1/widgets" in text
    assert "/etc/transagent/widgets.yaml" in text
    assert "curl -X POST" in text
    assert len(doc.tables) == 1


def test_mixed_probe_fixture_covers_go_no_go_inputs(okapi_probe_docx_path):
    doc = Document(okapi_probe_docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert len(doc.inline_shapes) == 1
    assert len(doc.tables) == 1
    assert "Kubernetes Deployment 文档" in text
    assert "D2_BODY_MARK_A17C" in text
    assert "D2_TABLE_MARK_E42D" in "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "D2_HEADER_MARK_6F8B" in doc.sections[0].header.paragraphs[0].text
    assert "D2_FOOTER_MARK_91C2" in doc.sections[0].footer.paragraphs[0].text
