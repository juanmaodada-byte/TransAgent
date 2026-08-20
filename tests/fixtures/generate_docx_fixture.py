"""
Deterministic fixture generator for TransAgent document pipeline D1.

Usage:
    python3 tests/fixtures/generate_docx_fixture.py
"""

from __future__ import annotations

import os
import shutil
import struct
import subprocess
import tempfile
import zipfile
from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
EXCEPTION_DIR = os.path.join(HERE, "format_exceptions")

PNG_DATA = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
    b"\x00\x00\x00\x0cIDATx\x9cc\x62\x60\x60\x00\x00\x00\x04\x00\x01\xa2\x17\xa5\x16"
    b"\x00\x00\x00\x00IEND\xaeB`\x82"
)
FIXED_DT = datetime(2026, 8, 13, 0, 0, 0, tzinfo=timezone.utc)
FIXED_ZIP_DT = (2026, 8, 13, 0, 0, 0)
BUNDLED_SOFFICE = os.environ.get("SOFFICE_PATH") or shutil.which("soffice") or ""


def save_docx_deterministic(doc, path):
    props = doc.core_properties
    props.author = "TransAgent"
    props.last_modified_by = "TransAgent"
    props.created = FIXED_DT
    props.modified = FIXED_DT
    doc.save(path)
    _normalize_zip(path)


def _normalize_zip(path):
    tmp_path = path + ".tmp"
    with zipfile.ZipFile(path, "r") as src, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for name in sorted(src.namelist()):
            info = zipfile.ZipInfo(name, FIXED_ZIP_DT)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            dst.writestr(info, src.read(name))
    os.replace(tmp_path, path)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = paragraph.add_run(text)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def build_plain_docx():
    doc = Document()
    doc.add_heading("Plain Technical Guide", level=1)
    p = doc.add_paragraph()
    p.add_run("This paragraph contains ")
    p.add_run("bold").bold = True
    p.add_run(" and ")
    p.add_run("italic").italic = True
    p.add_run(" runs inside the same paragraph.")
    doc.add_paragraph("Install the service with the following checklist:")
    for item in ["Prepare credentials", "Apply the manifest", "Verify the rollout"]:
        doc.add_paragraph(item, style="List Bullet")
    save_docx_deterministic(doc, os.path.join(HERE, "plain_technical.docx"))


def build_api_docx():
    doc = Document()
    doc.add_heading("Widget API Reference", level=1)
    doc.add_paragraph("Endpoint: https://api.example.com/v1/widgets")
    doc.add_paragraph("Config path: /etc/transagent/widgets.yaml")
    add_code_block(doc, "curl -X POST https://api.example.com/v1/widgets -d '{\"name\":\"demo\"}'")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Parameter", "Type", "Required", "Description"]
    rows = [
        ["name", "string", "yes", "Widget display name"],
        ["role", "string", "no", "One of: admin, developer, viewer"],
        ["limit", "integer", "no", "Maximum results per page"],
    ]
    for col, text in enumerate(headers):
        table.rows[0].cells[col].text = text
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = text
    save_docx_deterministic(doc, os.path.join(HERE, "api_reference.docx"))


def build_mixed_docx():
    doc = Document()

    section = doc.sections[0]
    section.header.paragraphs[0].text = "D2_HEADER_MARK_6F8B"
    section.footer.paragraphs[0].text = "D2_FOOTER_MARK_91C2"

    doc.add_heading("云原生应用部署指南", level=1)
    doc.add_paragraph("D2_BODY_MARK_A17C")
    doc.add_paragraph(
        "本文档介绍如何在 Kubernetes 集群中部署一个典型的云原生应用。"
        "该应用包含 Web 前端、API 服务和 Redis 缓存层。"
    )
    doc.add_heading("1. 环境准备", level=2)
    doc.add_paragraph("在开始部署之前，请确保以下工具已安装并正确配置：")
    for item in [
        "kubectl 命令行工具（版本 v1.28.0 或更高）",
        "Helm v3.12+ 包管理器",
        "Docker Desktop 或 Rancher Desktop",
        "访问集群的 kubeconfig 配置文件",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.1 验证集群连接", level=3)
    doc.add_paragraph("运行以下命令验证集群连接状态：")
    add_code_block(doc, "$ kubectl cluster-info\n$ kubectl get nodes")

    doc.add_heading("2. 资源配置表", level=2)
    doc.add_paragraph("各服务的资源需求如下：")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).merge(table.cell(0, 1)).text = "服务资源"
    table.cell(0, 2).text = "内存请求"
    table.cell(0, 3).text = "副本数"
    table.cell(1, 0).text = "Web 前端"
    table.cell(1, 1).text = "100m"
    table.cell(1, 2).text = "128Mi"
    table.cell(1, 3).text = "2"
    table.cell(2, 0).merge(table.cell(3, 0)).text = "后端服务"
    table.cell(2, 1).text = "250m"
    table.cell(2, 2).text = "256Mi"
    table.cell(2, 3).text = "3"
    table.cell(3, 1).text = "50m"
    table.cell(3, 2).text = "64Mi"
    table.cell(3, 3).text = "1"
    table.cell(3, 3).paragraphs[0].add_run(" D2_TABLE_MARK_E42D")

    doc.add_paragraph("请根据实际负载情况调整上述资源配置。")
    doc.add_heading("3. 部署清单示例", level=2)
    add_code_block(doc, "apiVersion: apps/v1\nkind: Deployment\nmetadata:\n  name: web-frontend")

    doc.add_heading("4. 架构示意图", level=2)
    doc.add_picture(BytesIO(PNG_DATA), width=Inches(2))
    doc.add_paragraph("[图1] 系统架构示意图")

    p = doc.add_paragraph("更多信息请参考 ")
    add_hyperlink(p, "Kubernetes Deployment 文档", "https://kubernetes.io/docs/concepts/workloads/controllers/deployment/")
    p.add_run("。")

    doc.add_heading("5. 部署流程", level=2)
    add_code_block(doc, "```mermaid\ngraph TD\nA[开始部署] --> B{检查集群状态}\n```")
    doc.add_heading("6. 验证与监控", level=2)
    doc.add_paragraph("配置文件位于 /etc/kubernetes/manifests/ 目录下。")

    mixed_path = os.path.join(HERE, "cloud_native_mixed.docx")
    save_docx_deterministic(doc, mixed_path)
    save_docx_deterministic(doc, os.path.join(HERE, "okapi_probe_mixed.docx"))


def build_exception_fixtures():
    os.makedirs(EXCEPTION_DIR, exist_ok=True)
    with open(os.path.join(EXCEPTION_DIR, "pdf_named_docx.docx"), "wb") as f:
        f.write(_minimal_pdf_bytes())
    with open(os.path.join(EXCEPTION_DIR, "corrupt.docx"), "wb") as f:
        f.write(b"this is not a valid ZIP file")
    with zipfile.ZipFile(os.path.join(EXCEPTION_DIR, "plain_zip.docx"), "w") as zf:
        zf.writestr("readme.txt", "ordinary zip, not docx")
    _normalize_zip(os.path.join(EXCEPTION_DIR, "plain_zip.docx"))
    with open(os.path.join(EXCEPTION_DIR, "unsupported.bin"), "wb") as f:
        f.write(b"\x00\x01unsupported")
    with open(os.path.join(EXCEPTION_DIR, "minimal.pdf"), "wb") as f:
        f.write(_minimal_pdf_bytes())
    with open(os.path.join(EXCEPTION_DIR, "fake_ole.doc"), "wb") as f:
        f.write(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 504)
    _write_minimal_ole2(
        os.path.join(EXCEPTION_DIR, "minimal_word.doc"),
        stream_name="WordDocument",
        extension_stream_name="1Table",
    )
    _write_minimal_ole2(
        os.path.join(EXCEPTION_DIR, "workbook_named_doc.doc"),
        stream_name="Workbook",
        extension_stream_name="\x05SummaryInformation",
    )
    _write_minimal_ole2(
        os.path.join(EXCEPTION_DIR, "encrypted_ole.doc"),
        stream_name="EncryptedPackage",
        extension_stream_name="\x06DataSpaces",
    )


def resolve_fixture_soffice() -> str | None:
    """Resolve an executable LibreOffice binary for optional fixture generation."""
    candidates = [
        os.environ.get("SOFFICE_PATH"),
        BUNDLED_SOFFICE,
        shutil.which("soffice"),
    ]
    for candidate in candidates:
        if not candidate:
            continue
        path = os.path.abspath(os.path.expanduser(candidate))
        if os.path.isfile(path) and os.access(path, os.X_OK):
            return path
    return None


def build_real_doc_fixture():
    """Generate a real Word 97-2003 DOC fixture from the mixed DOCX fixture."""
    soffice = resolve_fixture_soffice()
    if not soffice:
        print("Skipped real DOC fixture: LibreOffice not available")
        return
    source_docx = os.path.join(HERE, "okapi_probe_mixed.docx")
    output_doc = os.path.join(HERE, "real_word_97_mixed.doc")
    with tempfile.TemporaryDirectory(prefix="transagent-doc-fixture-") as tmp:
        profile = os.path.join(tmp, "lo-profile")
        outdir = os.path.join(tmp, "out")
        os.makedirs(profile)
        os.makedirs(outdir)
        cmd = [
            soffice,
            f"-env:UserInstallation=file://{profile}",
            "--headless",
            "--convert-to",
            "doc:MS Word 97",
            "--outdir",
            outdir,
            source_docx,
        ]
        result = subprocess.run(
            cmd,
            text=True,
            capture_output=True,
            timeout=120,
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError("LibreOffice failed to generate real Word 97-2003 DOC fixture")
        generated = os.path.join(outdir, "okapi_probe_mixed.doc")
        if not os.path.exists(generated) or os.path.getsize(generated) == 0:
            raise RuntimeError("LibreOffice did not create real Word 97-2003 DOC fixture")
        shutil.copy2(generated, output_doc)


def _minimal_pdf_bytes():
    return (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] >> endobj\n"
        b"trailer << /Root 1 0 R >>\n%%EOF\n"
    )


def _write_minimal_ole2(path, stream_name, extension_stream_name):
    """Write a small deterministic CFB/OLE2 file with sparse directory indexes.

    The file is only a format-detection fixture, not an Office-openable document.
    Directory index 1 is intentionally empty and the real streams start at
    original indexes 2/3, which catches parsers that compress directory entries.
    """
    sector_size = 512
    free = 0xFFFFFFFF
    end = 0xFFFFFFFE
    fat_sector = 0xFFFFFFFD
    no_stream = 0xFFFFFFFF

    header = bytearray(sector_size)
    header[:8] = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    header[24:26] = struct.pack("<H", 0x003E)
    header[26:28] = struct.pack("<H", 0x0003)
    header[28:30] = struct.pack("<H", 0xFFFE)
    header[30:32] = struct.pack("<H", 9)
    header[32:34] = struct.pack("<H", 6)
    header[40:44] = struct.pack("<I", 0)
    header[44:48] = struct.pack("<I", 1)
    header[48:52] = struct.pack("<I", 0)
    header[56:60] = struct.pack("<I", 4096)
    header[60:64] = struct.pack("<I", free)
    header[64:68] = struct.pack("<I", 0)
    header[68:72] = struct.pack("<I", free)
    header[72:76] = struct.pack("<I", 0)
    difat = [1] + [free] * 108
    header[76:512] = struct.pack("<109I", *difat)

    directory = bytearray(sector_size)
    _write_dir_entry(
        directory,
        0,
        "Root Entry",
        obj_type=5,
        child=2,
        start_sector=end,
        stream_size=0,
    )
    # Entry 1 deliberately left invalid/empty.
    _write_dir_entry(
        directory,
        2,
        stream_name,
        obj_type=2,
        right=3,
        start_sector=end,
        stream_size=0,
    )
    _write_dir_entry(
        directory,
        3,
        extension_stream_name,
        obj_type=2,
        start_sector=end,
        stream_size=0,
    )

    fat_values = [end, fat_sector] + [free] * 126
    fat = struct.pack("<128I", *fat_values)

    with open(path, "wb") as f:
        f.write(header)
        f.write(directory)
        f.write(fat)


def _write_dir_entry(
    directory,
    index,
    name,
    obj_type,
    left=0xFFFFFFFF,
    right=0xFFFFFFFF,
    child=0xFFFFFFFF,
    start_sector=0xFFFFFFFE,
    stream_size=0,
):
    offset = index * 128
    encoded_name = name.encode("utf-16le") + b"\x00\x00"
    if len(encoded_name) > 64:
        raise ValueError(f"OLE2 directory name too long: {name}")
    entry = bytearray(128)
    entry[:len(encoded_name)] = encoded_name
    entry[64:66] = struct.pack("<H", len(encoded_name))
    entry[66] = obj_type
    entry[67] = 1
    entry[68:72] = struct.pack("<I", left)
    entry[72:76] = struct.pack("<I", right)
    entry[76:80] = struct.pack("<I", child)
    entry[116:120] = struct.pack("<I", start_sector)
    entry[120:128] = struct.pack("<Q", stream_size)
    directory[offset:offset + 128] = entry


def build_document():
    build_plain_docx()
    build_api_docx()
    build_mixed_docx()
    build_exception_fixtures()
    build_real_doc_fixture()
    print(f"Generated fixtures under: {HERE}")


if __name__ == "__main__":
    build_document()
