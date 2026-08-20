"""D6 PDF-to-approximate-DOCX probe helpers.

This module is intentionally outside the production `extract_document()` path.
PyMuPDF and pdf2docx run only inside the fixed local `.runtime/pdf/` virtualenv
via worker subprocesses so system Python remains untouched.
"""

from __future__ import annotations

import hashlib
import importlib.metadata as metadata
import json
import os
import re
import shutil
import subprocess
import sys
import time
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
PDF_RUNTIME = ROOT / ".runtime" / "pdf" / "venv" / "bin" / "python"
PDF_INSPECT_WORKER = ROOT / "scripts" / "pdf_inspect_worker.py"
PDF_TO_DOCX_WORKER = ROOT / "scripts" / "pdf_to_docx_worker.py"
PDF_FIXTURE_GENERATOR = ROOT / "scripts" / "generate_pdf_fixtures.py"
MAX_INPUT_PDF_BYTES = 50 * 1024 * 1024
MAX_OUTPUT_DOCX_BYTES = 200 * 1024 * 1024
CONVERSION_TIMEOUT_SECONDS = 180
INSPECTION_TIMEOUT_SECONDS = 60
FALLBACK_WARNING = "pdf2docx conversion failed; delivered page-ordered text DOCX with layout and non-text content loss."
MIXED_TEXT_WARNING = "PDF contains pages without extractable text; image-only page content is not translated."


class PdfProbeError(ValueError):
    """Stable D6 PDF probe failure."""


def pdf_error(code: str, detail: str) -> PdfProbeError:
    return PdfProbeError(f"{code}: {detail}")


@dataclass
class CommandResult:
    cmd: list[str]
    returncode: int
    stdout: str
    stderr: str
    timed_out: bool = False


@dataclass
class PdfInspection:
    openable: bool
    encrypted: bool
    page_count: int = 0
    page_text_char_counts: list[int] = field(default_factory=list)
    text_pages: list[int] = field(default_factory=list)
    no_text_pages: list[int] = field(default_factory=list)
    total_text_chars: int = 0
    classification: str = ""
    error_code: str = ""
    warnings: list[str] = field(default_factory=list)
    pages_text: list[str] = field(default_factory=list)


@dataclass
class ConversionResult:
    engine: str
    input_pdf: str
    output_docx: str
    duration_seconds: float
    fallback_used: bool
    warnings: list[str]
    command: CommandResult | None = None


def ensure_pdf_runtime() -> dict:
    if not PDF_RUNTIME.exists() or not os.access(PDF_RUNTIME, os.X_OK):
        raise pdf_error("DOCUMENT_RUNTIME_UNAVAILABLE", f"fixed PDF runtime missing: {PDF_RUNTIME}")
    script = (
        "import importlib.metadata as md, sys, platform, json\n"
        "import fitz\n"
        "packages=['pdf2docx','PyMuPDF','python-docx','reportlab','pypdf']\n"
        "print(json.dumps({\n"
        " 'python': sys.version.split()[0],\n"
        " 'executable': sys.executable,\n"
        " 'platform': platform.platform(),\n"
        " 'machine': platform.machine(),\n"
        " 'packages': {name: md.version(name) for name in packages},\n"
        " 'fitz_version': fitz.version,\n"
        "}))\n"
    )
    result = subprocess.run([str(PDF_RUNTIME), "-c", script], text=True, capture_output=True, timeout=30, check=False)
    if result.returncode != 0:
        raise pdf_error("DOCUMENT_RUNTIME_UNAVAILABLE", (result.stderr or result.stdout).strip())
    text = result.stdout.strip().splitlines()[-1]
    info = json.loads(text)
    if info["packages"].get("pdf2docx") != "0.5.13":
        raise pdf_error("DOCUMENT_RUNTIME_UNAVAILABLE", f"pdf2docx version is {info['packages'].get('pdf2docx')}, expected 0.5.13")
    return info


def ensure_pdf_fixtures() -> Path:
    if not PDF_RUNTIME.exists():
        raise pdf_error("DOCUMENT_RUNTIME_UNAVAILABLE", "fixed PDF runtime is unavailable")
    result = subprocess.run(
        [str(PDF_RUNTIME), str(PDF_FIXTURE_GENERATOR)],
        text=True,
        capture_output=True,
        timeout=120,
        check=False,
    )
    if result.returncode != 0:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", (result.stderr or result.stdout).strip())
    return ROOT / "tests" / "fixtures" / "pdf"


def inspect_pdf(pdf_path: Path, include_text: bool = False) -> PdfInspection:
    if pdf_path.stat().st_size > MAX_INPUT_PDF_BYTES:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "PDF exceeds D6 probe input size limit")
    cmd = [str(PDF_RUNTIME), str(PDF_INSPECT_WORKER), "--input", str(pdf_path)]
    if include_text:
        cmd.append("--include-text")
    try:
        result = subprocess.run(cmd, text=True, capture_output=True, timeout=INSPECTION_TIMEOUT_SECONDS, check=False)
    except subprocess.TimeoutExpired as exc:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "PDF text-layer inspection timed out") from exc
    if result.returncode != 0:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", (result.stderr or result.stdout).strip())
    try:
        raw = json.loads(result.stdout.strip().splitlines()[-1])
    except (json.JSONDecodeError, IndexError) as exc:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "PDF inspection worker returned invalid JSON") from exc

    warnings: list[str] = []
    if raw.get("openable") and not raw.get("encrypted") and raw.get("text_pages") and raw.get("no_text_pages"):
        raw["classification"] = "mixed"
        warnings.append(MIXED_TEXT_WARNING)
    return PdfInspection(
        openable=bool(raw.get("openable")),
        encrypted=bool(raw.get("encrypted")),
        page_count=int(raw.get("page_count") or 0),
        page_text_char_counts=list(raw.get("page_text_char_counts") or []),
        text_pages=list(raw.get("text_pages") or []),
        no_text_pages=list(raw.get("no_text_pages") or []),
        total_text_chars=int(raw.get("total_text_chars") or 0),
        classification=str(raw.get("classification") or ""),
        error_code=str(raw.get("error_code") or ""),
        warnings=warnings,
        pages_text=list(raw.get("pages_text") or []),
    )


def require_convertible_pdf(pdf_path: Path) -> PdfInspection:
    inspection = inspect_pdf(pdf_path)
    if not inspection.openable or inspection.encrypted:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "PDF is encrypted or not parseable")
    if not inspection.text_pages:
        raise pdf_error("DOCUMENT_OCR_UNSUPPORTED", "PDF has no extractable text layer")
    return inspection


def convert_pdf_to_docx(pdf_path: Path, output_docx: Path, timeout: int = CONVERSION_TIMEOUT_SECONDS) -> ConversionResult:
    if pdf_path.resolve() == output_docx.resolve():
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "conversion would overwrite input PDF")
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    started = time.monotonic()
    cmd = [
        str(PDF_RUNTIME),
        str(PDF_TO_DOCX_WORKER),
        "--input",
        str(pdf_path),
        "--output",
        str(output_docx),
    ]
    timed_out = False
    try:
        result = subprocess.run(cmd, text=True, capture_output=True, timeout=timeout, check=False)
    except subprocess.TimeoutExpired as exc:
        timed_out = True
        result = subprocess.CompletedProcess(cmd, 124, exc.stdout or "", exc.stderr or "conversion timed out")
    command = CommandResult(cmd=cmd, returncode=result.returncode, stdout=result.stdout or "", stderr=result.stderr or "", timed_out=timed_out)
    duration = time.monotonic() - started
    if timed_out:
        raise pdf_error("DOCUMENT_CONVERSION_ERROR", "pdf2docx conversion timed out")
    if result.returncode != 0:
        raise pdf_error("DOCUMENT_CONVERSION_ERROR", (result.stderr or result.stdout).strip())
    validate_docx_package(output_docx)
    return ConversionResult(
        engine="pdf2docx 0.5.13",
        input_pdf=str(pdf_path),
        output_docx=str(output_docx),
        duration_seconds=duration,
        fallback_used=False,
        warnings=[],
        command=command,
    )


def fallback_text_docx(pdf_path: Path, output_docx: Path) -> ConversionResult:
    inspection = inspect_pdf(pdf_path, include_text=True)
    if not inspection.openable or inspection.encrypted:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "fallback requires parseable unencrypted PDF")
    if not inspection.text_pages:
        raise pdf_error("DOCUMENT_OCR_UNSUPPORTED", "fallback requires extractable text")
    started = time.monotonic()
    doc = Document()
    non_empty_paragraphs = 0
    for page_index, page_text in enumerate(inspection.pages_text):
        if page_index:
            doc.add_page_break()
        wrote_on_page = False
        for line in page_text.splitlines():
            if not line and not wrote_on_page:
                continue
            if not line and (not doc.paragraphs or not doc.paragraphs[-1].text):
                continue
            doc.add_paragraph(line)
            wrote_on_page = wrote_on_page or bool(line)
            if line:
                non_empty_paragraphs += 1
    if non_empty_paragraphs == 0:
        raise pdf_error("DOCUMENT_OCR_UNSUPPORTED", "fallback produced no text paragraphs")
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    validate_docx_package(output_docx)
    return ConversionResult(
        engine="pymupdf text -> python-docx",
        input_pdf=str(pdf_path),
        output_docx=str(output_docx),
        duration_seconds=time.monotonic() - started,
        fallback_used=True,
        warnings=[FALLBACK_WARNING],
    )


def validate_docx_package(docx_path: Path) -> dict:
    if not docx_path.exists():
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", f"DOCX output missing: {docx_path}")
    size = docx_path.stat().st_size
    if size <= 0:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "DOCX output is empty")
    if size > MAX_OUTPUT_DOCX_BYTES:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "DOCX output exceeds D6 size limit")
    if not zipfile.is_zipfile(docx_path):
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "DOCX output is not a ZIP package")
    xml_files: list[str] = []
    rels_files: list[str] = []
    try:
        with zipfile.ZipFile(docx_path) as zf:
            bad = zf.testzip()
            if bad:
                raise pdf_error("DOCUMENT_INTEGRITY_ERROR", f"DOCX ZIP CRC failed for member {bad}")
            names = set(zf.namelist())
            missing = {"[Content_Types].xml", "word/document.xml"} - names
            if missing:
                raise pdf_error("DOCUMENT_INTEGRITY_ERROR", f"DOCX missing required package parts: {sorted(missing)}")
            for name in sorted(names):
                if name.endswith(".xml"):
                    ET.fromstring(zf.read(name))
                    xml_files.append(name)
                elif name.endswith(".rels"):
                    ET.fromstring(zf.read(name))
                    rels_files.append(name)
                if name.endswith(".xml") or name.endswith(".rels"):
                    if b"[[TA_" in zf.read(name):
                        raise pdf_error("DOCUMENT_PLACEHOLDER_ERROR", "internal placeholder remained in DOCX")
    except zipfile.BadZipFile as exc:
        raise pdf_error("DOCUMENT_INTEGRITY_ERROR", "DOCX output is not a valid ZIP") from exc
    return {"size_bytes": size, "xml_file_count": len(xml_files), "rels_file_count": len(rels_files)}


def docx_contains_replacement_char(docx_path: Path) -> bool:
    with zipfile.ZipFile(docx_path) as zf:
        for name in zf.namelist():
            if name.endswith(".xml") and "\ufffd" in zf.read(name).decode("utf-8", errors="replace"):
                return True
    return False


def docx_has_residual_placeholders(docx_path: Path) -> bool:
    with zipfile.ZipFile(docx_path) as zf:
        return any(b"[[TA_" in zf.read(name) for name in zf.namelist() if name.endswith(".xml") or name.endswith(".rels"))


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def compare_snapshots(before: dict, after: dict) -> dict:
    diffs: dict = {}
    for key in sorted(set(before) | set(after)):
        if before.get(key) != after.get(key):
            diffs[key] = {"before": before.get(key), "after": after.get(key)}
    return diffs


def make_pseudo_blocks(blocks: list) -> list:
    from interface import DocumentBlock

    return [
        DocumentBlock(block_id=block.block_id, block_type=block.block_type, text=f"[D6-ZH] {block.text}")
        for block in blocks
    ]


def write_runtime_manifest(path: Path, runtime: dict) -> None:
    packages = runtime.get("packages", {})
    lines = [
        "# D6 PDF Runtime Manifest",
        "",
        "Status: fixed local PDF runtime verified.",
        "",
        "## Python",
        "",
        f"- Version: `{runtime.get('python', 'unknown')}`",
        f"- Platform: `{runtime.get('platform', 'unknown')}`",
        f"- Machine: `{runtime.get('machine', 'unknown')}`",
        f"- Local path: `.runtime/pdf/venv/bin/python`",
        f"- Absolute run path: `{runtime.get('executable', PDF_RUNTIME)}`",
        "",
        "## Packages",
        "",
        "| Package | Version | Source |",
        "|---|---:|---|",
    ]
    for package in ["pdf2docx", "PyMuPDF", "python-docx", "reportlab", "pypdf"]:
        lines.append(f"| `{package}` | `{packages.get(package, 'unknown')}` | PyPI via project-local venv install |")
    lines.extend([
        "",
        "## Boundaries",
        "",
        "- No system-level or Homebrew installation was used.",
        "- `pdf2docx` was not silently changed; required version is `0.5.13`.",
        "- PyMuPDF/pdf2docx are invoked through fixed worker subprocesses, not product imports.",
        "- Runtime paths are probe configuration, not final product dependencies.",
        "",
    ])
    path.write_text("\n".join(lines), encoding="utf-8")


def package_versions_for_current_python() -> dict:
    versions = {}
    for package in ("python-docx", "pytest"):
        try:
            versions[package] = metadata.version(package)
        except metadata.PackageNotFoundError:
            versions[package] = ""
    return versions
