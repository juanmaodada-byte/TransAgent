"""
预处理入口
==========
Vibe Coder B | v1.0 | 2026-08-06

一站式预处理：格式检测 → 格式转换 → 结构解析 → 文档分块。

使用：
    from transagent.backend.pipeline.preprocess import preprocess
    result = preprocess("/path/to/doc.docx")
"""

import os
import posixpath
import re
import struct
import tempfile
import zipfile
from pathlib import Path, PurePosixPath
from xml.etree import ElementTree as ET
from transagent.interface import (
    FormatResult, ConvertResult, PreprocessResult, PlaceholderMap, Chunk, FormatType
)
from transagent.backend.pipeline.structure_parser import parse_structure
from transagent.backend.pipeline.chunker import chunk_document
from transagent.backend.config import get_config


def preprocess(file_path: str) -> PreprocessResult:
    """
    一站式预处理入口。

    流程: 格式检测 → MD转换 → 结构解析 → 文档分块

    Args:
        file_path: 用户上传文件路径

    Returns:
        PreprocessResult（受保护MD + chunk列表 + 占位符映射表）
    """
    # Step 1: 格式检测
    fmt = detect_format(file_path)

    # Step 2: 格式→MD
    converted = convert_to_md(file_path, fmt.format_type)

    # Step 3: 结构解析（占位符保护）
    protected_md, pmap = parse_structure(converted.md_text)

    # Step 4: 文档分块
    chunks = chunk_document(protected_md)

    # 转换警告（DOC 归一化 / PDF 近似转换等）→ 透传到 PreprocessResult
    conversion_warnings = list(converted.metadata.get("conversion_warnings", []))

    return PreprocessResult(
        protected_md=protected_md,
        chunks=chunks,
        placeholder_map=pmap,
        token_estimate_total=sum(c.token_estimate for c in chunks),
        chunk_count=len(chunks),
        conversion_warnings=conversion_warnings,
    )


DOC_ERROR_MISMATCH = "DOCUMENT_FORMAT_MISMATCH"
DOC_ERROR_UNSUPPORTED = "DOCUMENT_UNSUPPORTED_FORMAT"
DOC_ERROR_INTEGRITY = "DOCUMENT_INTEGRITY_ERROR"

_OOXML_CONTENT_TYPES_NS = {
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types"
}
_DOCX_MAIN_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
}
_DOCX_MACRO_CONTENT_TYPES = {
    "application/vnd.ms-word.document.macroEnabled.main+xml",
    "application/vnd.ms-word.template.macroEnabledTemplate.main+xml",
}
_OOXML_ENCRYPTED_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.encryptedPackage",
}

_ZIP_MAX_MEMBERS = 1000
_ZIP_MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
_ZIP_MAX_COMPRESSION_RATIO = 100
_OLE2_HEADER = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


class DocumentFormatError(ValueError):
    """文档格式错误，消息以稳定错误码开头。"""


def _format_error(code: str, detail: str) -> DocumentFormatError:
    return DocumentFormatError(f"{code}: {detail}")


def detect_format(file_path: str) -> FormatResult:
    """检测文件真实格式，不仅依赖扩展名。"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    size = os.path.getsize(file_path)

    ext_map = {
        ".md": FormatType.MARKDOWN,
        ".markdown": FormatType.MARKDOWN,
        ".docx": FormatType.DOCX,
        ".doc": FormatType.DOC,
        ".pdf": FormatType.PDF,
        ".txt": FormatType.TEXT,
    }
    expected = ext_map.get(ext)
    if expected is None:
        raise _format_error(
            DOC_ERROR_UNSUPPORTED,
            f"unsupported extension: {ext or '(none)'}"
        )

    cfg = get_config().app
    if size > cfg.max_file_size_bytes:
        raise _format_error(
            DOC_ERROR_INTEGRITY,
            f"file too large: {size / 1024 / 1024:.1f}MB "
            f"(limit {cfg.max_file_size_bytes / 1024 / 1024:.0f}MB)"
        )

    detected, metadata = _detect_real_container(file_path)
    if detected != expected.value:
        raise _format_error(
            DOC_ERROR_MISMATCH,
            f"extension {ext or '(none)'} expects {expected.value}, "
            f"but content is {detected}"
        )

    page_count = None
    if detected == FormatType.PDF.value:
        page_count = metadata.get("page_count")

    return FormatResult(
        format_type=detected,
        mime_type=_mime_for_format(detected),
        size_bytes=size,
        page_count=page_count,
        metadata={
            "extension": ext,
            "filename": os.path.basename(file_path),
            **metadata,
        },
    )


def _detect_real_container(file_path: str) -> tuple[str, dict]:
    with open(file_path, "rb") as f:
        head = f.read(8192)

    if head.startswith(b"%PDF-"):
        return _validate_pdf(file_path)
    if head.startswith(_OLE2_HEADER):
        return _validate_doc_ole(file_path)
    if head.startswith(b"PK\x03\x04") or zipfile.is_zipfile(file_path):
        return _validate_docx_zip(file_path)
    if os.path.splitext(file_path)[1].lower() == ".docx":
        raise _format_error(DOC_ERROR_INTEGRITY, "DOCX extension but ZIP header is missing")

    ext = os.path.splitext(file_path)[1].lower()
    if ext in (".md", ".markdown"):
        _validate_utf8_text(file_path)
        return FormatType.MARKDOWN.value, {"detector": "utf8-text"}
    if ext == ".txt":
        _validate_utf8_text(file_path)
        return FormatType.TEXT.value, {"detector": "utf8-text"}

    raise _format_error(DOC_ERROR_UNSUPPORTED, "content signature is unsupported")


def _validate_docx_zip(file_path: str) -> tuple[str, dict]:
    try:
        with zipfile.ZipFile(file_path) as zf:
            members = zf.infolist()
            if len(members) > _ZIP_MAX_MEMBERS:
                raise _format_error(DOC_ERROR_INTEGRITY, "zip has too many members")
            names = {info.filename for info in members}
            total_size = 0
            for info in members:
                _validate_zip_member(info)
                total_size += info.file_size
                if total_size > _ZIP_MAX_UNCOMPRESSED_BYTES:
                    raise _format_error(DOC_ERROR_INTEGRITY, "zip expands too large")

            required = {"[Content_Types].xml", "word/document.xml"}
            missing = sorted(required - names)
            if missing:
                raise _format_error(
                    DOC_ERROR_UNSUPPORTED,
                    f"zip is not a DOCX package; missing {', '.join(missing)}"
                )

            content_types = zf.read("[Content_Types].xml")
            _validate_docx_content_types(content_types)
            try:
                ET.fromstring(zf.read("word/document.xml"))
            except ET.ParseError as e:
                raise _format_error(
                    DOC_ERROR_INTEGRITY,
                    f"word/document.xml is not parseable XML: {e}"
                ) from e
    except zipfile.BadZipFile as e:
        raise _format_error(DOC_ERROR_INTEGRITY, "corrupt ZIP/DOCX container") from e
    except RuntimeError as e:
        raise _format_error(DOC_ERROR_INTEGRITY, f"unreadable ZIP member: {e}") from e

    return FormatType.DOCX.value, {"detector": "ooxml-wordprocessingml"}


def _validate_zip_member(info: zipfile.ZipInfo) -> None:
    name = info.filename
    normalized = posixpath.normpath(name)
    path = PurePosixPath(name)
    if name.startswith("/") or normalized.startswith("../") or ".." in path.parts:
        raise _format_error(DOC_ERROR_INTEGRITY, f"unsafe ZIP member path: {name}")
    if info.compress_size > 0 and info.file_size / info.compress_size > _ZIP_MAX_COMPRESSION_RATIO:
        raise _format_error(DOC_ERROR_INTEGRITY, "suspicious ZIP compression ratio")
    if info.compress_size == 0 and info.file_size > 0:
        raise _format_error(DOC_ERROR_INTEGRITY, "suspicious empty compressed member")


def _validate_docx_content_types(raw_xml: bytes) -> None:
    try:
        root = ET.fromstring(raw_xml)
    except ET.ParseError as e:
        raise _format_error(
            DOC_ERROR_INTEGRITY,
            f"[Content_Types].xml is not parseable XML: {e}"
        ) from e

    main_types = []
    for override in root.findall("ct:Override", _OOXML_CONTENT_TYPES_NS):
        part_name = override.get("PartName", "")
        content_type = override.get("ContentType", "")
        if content_type in _DOCX_MACRO_CONTENT_TYPES:
            raise _format_error(DOC_ERROR_UNSUPPORTED, "macro-enabled DOCX/DOCM is unsupported")
        if content_type in _OOXML_ENCRYPTED_CONTENT_TYPES:
            raise _format_error(DOC_ERROR_INTEGRITY, "encrypted OOXML package is unsupported")
        if part_name == "/word/document.xml":
            main_types.append(content_type)

    if not main_types:
        raise _format_error(DOC_ERROR_UNSUPPORTED, "DOCX main document content type is missing")
    if not any(ct in _DOCX_MAIN_CONTENT_TYPES for ct in main_types):
        raise _format_error(
            DOC_ERROR_UNSUPPORTED,
            f"not a WordprocessingML main document: {main_types[0]}"
        )


def _validate_doc_ole(file_path: str) -> tuple[str, dict]:
    try:
        streams = _list_ole_streams(file_path)
    except Exception as e:
        raise _format_error(DOC_ERROR_INTEGRITY, f"unreadable OLE2 container: {e}") from e

    stream_names = {"/".join(parts) for parts in streams}
    leaf_names = {parts[-1] for parts in streams if parts}
    if "EncryptedPackage" in leaf_names:
        raise _format_error(DOC_ERROR_INTEGRITY, "encrypted OLE2 document is unsupported")
    if "WordDocument" not in leaf_names:
        raise _format_error(
            DOC_ERROR_UNSUPPORTED,
            "OLE2 container is not confirmed as Word DOC"
        )
    return FormatType.DOC.value, {
        "detector": "ole2-worddocument",
        "ole_streams": sorted(stream_names),
    }


def _list_ole_streams(file_path: str) -> list[list[str]]:
    with open(file_path, "rb") as f:
        header = f.read(512)
        if len(header) < 512 or not header.startswith(_OLE2_HEADER):
            raise ValueError("missing OLE2 header")

        sector_shift = struct.unpack_from("<H", header, 30)[0]
        mini_sector_shift = struct.unpack_from("<H", header, 32)[0]
        if sector_shift not in (9, 12) or mini_sector_shift != 6:
            raise ValueError("unsupported OLE2 sector size")
        sector_size = 1 << sector_shift
        num_fat_sectors = struct.unpack_from("<I", header, 44)[0]
        first_dir_sector = struct.unpack_from("<I", header, 48)[0]
        difat = list(struct.unpack_from("<109I", header, 76))
        fat_sector_ids = [sid for sid in difat if sid != 0xFFFFFFFF][:num_fat_sectors]
        if not fat_sector_ids:
            raise ValueError("missing OLE2 FAT")

        def read_sector(sector_id: int) -> bytes:
            if sector_id in (0xFFFFFFFE, 0xFFFFFFFF):
                raise ValueError("invalid sector id")
            f.seek((sector_id + 1) * sector_size)
            data = f.read(sector_size)
            if len(data) != sector_size:
                raise ValueError("truncated OLE2 sector")
            return data

        fat = []
        for sid in fat_sector_ids:
            sector = read_sector(sid)
            fat.extend(struct.unpack("<" + "I" * (sector_size // 4), sector))

        def read_chain(start_sid: int, max_sectors: int = 4096) -> bytes:
            if start_sid in (0xFFFFFFFE, 0xFFFFFFFF):
                return b""
            out = bytearray()
            sid = start_sid
            seen = set()
            while sid not in (0xFFFFFFFE, 0xFFFFFFFF):
                if sid in seen or sid >= len(fat) or len(seen) >= max_sectors:
                    raise ValueError("invalid OLE2 FAT chain")
                seen.add(sid)
                out.extend(read_sector(sid))
                sid = fat[sid]
            return bytes(out)

        directory = read_chain(first_dir_sector)

    entries_by_index = {}
    for index, offset in enumerate(range(0, len(directory), 128)):
        entry = directory[offset:offset + 128]
        if len(entry) < 128:
            continue
        name_len = struct.unpack_from("<H", entry, 64)[0]
        obj_type = entry[66]
        if obj_type not in (1, 2, 5) or name_len < 2:
            continue
        raw_name = entry[:name_len - 2]
        try:
            name = raw_name.decode("utf-16le")
        except UnicodeDecodeError:
            continue
        left = struct.unpack_from("<I", entry, 68)[0]
        right = struct.unpack_from("<I", entry, 72)[0]
        child = struct.unpack_from("<I", entry, 76)[0]
        entries_by_index[index] = {
            "name": name,
            "type": obj_type,
            "left": left,
            "right": right,
            "child": child,
        }

    streams = []

    def walk(index: int, prefix: list[str], depth: int = 0) -> None:
        if index == 0xFFFFFFFF or depth > len(entries_by_index):
            return
        item = entries_by_index.get(index)
        if item is None:
            return
        walk(item["left"], prefix, depth + 1)
        if item["type"] == 1:
            child_prefix = prefix + [item["name"]]
            walk(item["child"], child_prefix, depth + 1)
        elif item["type"] == 2:
            streams.append(prefix + [item["name"]])
        walk(item["right"], prefix, depth + 1)

    root = entries_by_index.get(0)
    if root is None or root["type"] != 5:
        raise ValueError("missing OLE2 root storage")
    walk(root["child"], [])
    return streams


def _validate_pdf(file_path: str) -> tuple[str, dict]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        page_count = _validate_pdf_minimal(file_path)
        return FormatType.PDF.value, {
            "detector": "pdf-header-stdlib",
            "page_count": page_count,
            "validation_warning": "PyMuPDF unavailable; used minimal PDF structure validation",
        }
    try:
        doc = fitz.open(file_path)
        page_count = doc.page_count
        if page_count < 1:
            raise _format_error(DOC_ERROR_INTEGRITY, "PDF has no pages")
        doc.close()
    except DocumentFormatError:
        raise
    except Exception as e:
        raise _format_error(DOC_ERROR_INTEGRITY, f"PDF parser could not open file: {e}") from e
    return FormatType.PDF.value, {"detector": "pdf-header-pymupdf", "page_count": page_count}


def _validate_pdf_minimal(file_path: str) -> int:
    with open(file_path, "rb") as f:
        data = f.read(1024 * 1024 + 1)
    if len(data) > 1024 * 1024:
        raise _format_error(DOC_ERROR_INTEGRITY, "PDF fallback validation limit exceeded")
    if not data.startswith(b"%PDF-") or b"%%EOF" not in data:
        raise _format_error(DOC_ERROR_INTEGRITY, "invalid minimal PDF structure")
    page_markers = len(re.findall(rb"/Type\s*/Page\b", data))
    if page_markers < 1:
        raise _format_error(DOC_ERROR_INTEGRITY, "PDF has no detectable page object")
    return page_markers


def _validate_utf8_text(file_path: str) -> None:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            f.read()
    except UnicodeDecodeError as e:
        raise _format_error(DOC_ERROR_INTEGRITY, f"UTF-8 text decode failed: {e}") from e


def _mime_for_format(format_type: str) -> str:
    return {
        FormatType.MARKDOWN.value: "text/markdown",
        FormatType.TEXT.value: "text/plain",
        FormatType.DOCX.value: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        FormatType.DOC.value: "application/msword",
        FormatType.PDF.value: "application/pdf",
    }.get(format_type, "application/octet-stream")


def convert_to_md(file_path: str, format_type: str) -> ConvertResult:
    """
    任何格式→MD转换。

    P0支持: md(直接通过), txt(直接通过), docx(文本提取)
    P1支持: pdf(文本+嵌入图提取)
    """
    if format_type in (FormatType.MARKDOWN.value, FormatType.TEXT.value):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                md_text = f.read()
        except UnicodeDecodeError as e:
            raise ValueError(
                f"DOCUMENT_INTEGRITY_ERROR: 文件编码无效，无法以 UTF-8 解码。"
                f"{os.path.basename(file_path)} 可能使用了非 UTF-8 编码（如 GBK、Latin-1）。"
                f"请将文件转换为 UTF-8 编码后重试。"
                f"原始错误: {type(e).__name__}: {e}"
            ) from e
        return ConvertResult(
            md_text=md_text,
            assets_dir="",
            image_count=0,
            metadata={"converter": "passthrough"},
        )

    elif format_type == FormatType.DOCX.value:
        return _convert_docx(file_path)

    elif format_type == FormatType.DOC.value:
        return _convert_doc(file_path)

    elif format_type == FormatType.PDF.value:
        # pdf2docx 优先（结构/阅读顺序更好），失败回退 PyMuPDF 抽文本。
        result = _convert_pdf_via_pdf2docx(file_path)
        if result is not None:
            return result
        return _convert_pdf(file_path)

    else:
        raise ValueError(f"不支持的格式: {format_type}（支持 md/txt/docx/doc/pdf）")


def _convert_docx(file_path: str) -> ConvertResult:
    """docx→MD转换（P0核心功能）"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx未安装。pip install python-docx")

    cfg = get_config().app
    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(
            f"DOCUMENT_INTEGRITY_ERROR: 无法打开或解析 DOCX 文件。"
            f"{os.path.basename(file_path)} 可能已损坏或不是有效的 DOCX 格式。"
            f"原始错误: {type(e).__name__}: {e}"
        ) from e
    md_lines: list[str] = []
    image_count = 0
    assets_dir = os.path.join(cfg.workspace_dir, "assets")
    os.makedirs(assets_dir, exist_ok=True)

    for para in doc.paragraphs:
        # 检测标题
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 2
            md_lines.append(f"{'#' * level} {para.text}")
        else:
            md_lines.append(para.text)

    # 表格→MD管道表
    for table in doc.tables:
        md_lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " ") for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                md_lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        md_lines.append("")

    # 嵌入图片提取
    # 遍历 rels 提取图片
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            img_ext = os.path.splitext(rel.target_ref)[1]
            img_path = os.path.join(assets_dir, f"img_{image_count:02d}{img_ext}")
            with open(img_path, "wb") as f:
                f.write(rel.target_part.blob)

    return ConvertResult(
        md_text="\n\n".join(md_lines),
        assets_dir=assets_dir,
        image_count=image_count,
        metadata={"converter": "python-docx", "paragraphs": len(doc.paragraphs)},
    )


def _convert_doc(file_path: str) -> ConvertResult:
    """DOC→MD转换：先经 LibreOffice 归一化为 DOCX，再走 DOCX→MD。

    懒加载 doc_normalizer（其顶层 import detect_format 来自本模块，顶层导入会循环依赖）。
    LibreOffice 缺失时 normalize_doc_to_docx 抛 DOCUMENT_RUNTIME_UNAVAILABLE，自然透传。
    """
    from transagent.backend.pipeline.doc_normalizer import normalize_doc_to_docx

    with tempfile.TemporaryDirectory(prefix="transagent_doc_") as td:
        normalized = normalize_doc_to_docx(Path(file_path), Path(td))
        result = _convert_docx(str(normalized.path))
        warnings = list(normalized.warnings)
        result.metadata["conversion_warnings"] = warnings
        result.metadata["converter"] = "LibreOffice+python-docx"
        return result


def _convert_pdf_via_pdf2docx(file_path: str) -> ConvertResult | None:
    """PDF→MD（pdf2docx 优先）：进程内 pdf2docx 转 DOCX → 复用 _convert_docx。

    任何失败（未安装/转换异常）返回 None，由调用方回退 PyMuPDF 抽文本。
    """
    try:
        from pdf2docx import Converter
    except ImportError:
        return None

    try:
        with tempfile.TemporaryDirectory(prefix="transagent_pdf_") as td:
            docx_path = os.path.join(td, "converted.docx")
            cv = Converter(file_path)
            try:
                cv.convert(docx_path, start=0, end=None)
            finally:
                cv.close()
            result = _convert_docx(docx_path)
    except Exception:
        return None

    # 扫描件：pdf2docx 可能"成功"产出空 DOCX（如空表格/空段落）→ 需在此拦截
    # 判据：MD 中无任何字母/数字字符即视为无文本层（管道符 | 等不算）
    if not any(c.isalnum() for c in result.md_text):
        raise ValueError(
            "DOCUMENT_OCR_UNSUPPORTED: PDF 无提取文本层（扫描件暂不支持 OCR，请上传文本型 PDF）"
        )

    result.metadata["conversion_warnings"] = [
        "PDF was converted to DOCX approximately; layout, reading order, tables, and images may differ."
    ]
    result.metadata["converter"] = "pdf2docx+python-docx"
    result.md_text = _apply_pdf_headings(result.md_text, _detect_pdf_headings(file_path))
    return result


# ── PDF 标题层级检测（启发式·基于字体度量）────────────────────────
# PyMuPDF get_text() 只提取文字，丢失样式 → PDF 转 MD 没有 # 标题。
# 用 dict 模式读取每行字号/加粗，按启发式识别标题并映射到 MD 的 #/##/###。
# 规则：
#   字号 >= 正文×1.2  → 标题（比值 1.6+ → #，1.35+ → ##，否则 ###）
#   加粗 + 短行(≤60字) + 非元数据行 → ###（正文里加粗的小标题）
# 两个 PDF 转换路径（pdf2docx / PyMuPDF 回退）都套用。

_PDF_BODY_BOLD_FLAG = 16  # PyMuPDF text flags bit 4 = bold


def _norm_heading(text: str) -> str:
    """规整标题/行文本用于匹配：小写、折叠空白、去首尾标点。"""
    return re.sub(r"\s+", " ", text.strip().strip("*·-•")).lower()


def _looks_like_pdf_metadata(text: str) -> bool:
    """排除明显的元数据/列表行（不是标题）。"""
    lower = re.sub(r"^[*·•\-_\s]+", "", text).lower()
    for kw in ("received:", "accepted:", "published:", "corresponding",
               "email:", "issn:", "doi:", "keywords:", "abstract", "refere",
               "volume", "copyright", "open access"):
        if lower.startswith(kw):
            return True
    return False


def _detect_pdf_headings(file_path: str) -> list[tuple[str, int]]:
    """从 PDF 用字体度量识别标题行，返回 [(标题文本, 层级)]。

    启发式（宁可漏报，不要误报）：
      - 字号规则：每页最大字号 且 >= 正文×1.25 → #/##/###（按比值）——命中论文/章节大标题
      - 加粗规则：整行加粗 + 短行(≤60字) + ≥2词 + 非元数据 → ###（正文里加粗的小节标题）
      - 元数据/页眉（ISSN/Received/期刊名等）在两类规则前统一排除
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return []
    try:
        doc = fitz.open(file_path)
        body_sizes: list[float] = []
        page_lines: list[list[tuple[str, float, bool]]] = []
        for page in doc:
            lines_info: list[tuple[str, float, bool]] = []
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    max_size = max((s.get("size", 0) for s in spans), default=0.0)
                    all_bold = all(bool(s.get("flags", 0) & _PDF_BODY_BOLD_FLAG) for s in spans)
                    if text:
                        lines_info.append((text, max_size, all_bold))
                        if max_size > 0:
                            body_sizes.append(round(max_size, 1))
            page_lines.append(lines_info)
        doc.close()
    except Exception:
        return []
    if not body_sizes:
        return []

    body_size = max(set(body_sizes), key=body_sizes.count)
    raw: list[tuple[str, int, float]] = []  # (text, level, font_size)；font_size=0 表示加粗小节规则
    for lines_info in page_lines:
        for text, max_size, all_bold in lines_info:
            if len(text) < 2 or _looks_like_pdf_metadata(text):
                continue
            # 标题规则：字号显著大于正文（排除"加粗短行"——期刊名/栏目名常为加粗大字号）
            if max_size >= body_size * 1.4 and not (all_bold and len(text) <= 60) \
                    and len(text.split()) >= 3:
                ratio = max_size / body_size
                level = 1 if ratio >= 1.6 else (2 if ratio >= 1.35 else 3)
                raw.append((text, level, max_size))
                continue
            # 小节规则：加粗 + 字号≈正文（0.95~1.15 倍）+ 短行 → ###
            if all_bold and body_size * 0.95 <= max_size <= body_size * 1.15 \
                    and len(text) <= 60:
                raw.append((text, 3, 0.0))

    # 合并相邻"同字号"标题行（如长标题在 PDF 里换行成两行）
    headings: list[tuple[str, int]] = []
    last_size: float = 0.0
    for text, level, font_size in raw:
        if (headings and headings[-1][1] == level
                and font_size > 0 and last_size > 0
                and min(font_size, last_size) / max(font_size, last_size) >= 0.9):
            prev_text, prev_level = headings[-1]
            headings[-1] = (prev_text.rstrip() + " " + text, prev_level)
            last_size = max(font_size, last_size)  # 保持合并后标题的字号
        else:
            headings.append((text, level))
            last_size = font_size
    return headings


def _apply_pdf_headings(md_text: str, headings: list[tuple[str, int]]) -> str:
    """把检测到的标题映射到转换后 MD 的对应行，加上 #/##/###。"""
    if not headings or not md_text:
        return md_text
    sigs = sorted([(orig, _norm_heading(orig), lv) for orig, lv in headings],
                  key=lambda x: -len(x[1]))
    used = [False] * len(sigs)
    out: list[str] = []
    for line in md_text.split("\n"):
        norm = _norm_heading(line)
        applied = False
        if norm and not line.startswith("#"):
            for i, (orig, sig, lv) in enumerate(sigs):
                if used[i] or not sig:
                    continue
                if norm == sig:
                    out.append("#" * lv + " " + line)
                    used[i] = True
                    applied = True
                    break
                # 行以标题开头且尾部很短（如"标题 + 作者名"挤在同一行）→ 拆成标题行 + 尾部行
                if len(sig) >= 6 and norm.startswith(sig) and len(norm) - len(sig) <= 24:
                    idx = line.lower().find(orig.lower())
                    if idx == 0 and len(orig) < len(line):
                        out.append("#" * lv + " " + line[:len(orig)].rstrip())
                        tail = line[len(orig):].strip()
                        if tail:
                            out.append(tail)
                    else:
                        out.append("#" * lv + " " + line)
                    used[i] = True
                    applied = True
                    break
        if not applied:
            out.append(line)
    return "\n".join(out)


def _convert_pdf(file_path: str) -> ConvertResult:
    """PDF→MD转换（PyMuPDF 抽文本·pdf2docx 回退路径）"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF未安装。pip install PyMuPDF")

    doc = fitz.open(file_path)
    md_lines: list[str] = []

    for page in doc:
        text = page.get_text()
        md_lines.append(text)

    doc.close()
    if not any(c.isalnum() for c in "".join(md_lines)):
        raise ValueError(
            "DOCUMENT_OCR_UNSUPPORTED: PDF 无提取文本层（扫描件暂不支持 OCR，请上传文本型 PDF）"
        )
    md_text = _apply_pdf_headings("\n\n".join(md_lines), _detect_pdf_headings(file_path))
    return ConvertResult(
        md_text=md_text,
        assets_dir="",
        image_count=0,
        metadata={"converter": "PyMuPDF", "pages": len(md_lines)},
    )
