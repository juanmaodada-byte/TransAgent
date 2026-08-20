"""
导出器
======
Vibe Coder B | v1.1 | 2026-08-16

职责：翻译后的MD文本按需导出为目标格式。
      支持：md（原样 UTF-8）、docx（python-docx反向重建）、HTML、双语对照（docx·源文/译文左右对照）

使用：
    from transagent.backend.pipeline.exporter import export_to_format
    output_path = export_to_format(md_text, "docx", assets_dir)
    output_path = export_to_format(md_text, "bilingual", assets_dir,
                                   aligned_pairs=pairs, placeholder_map=pmap)
"""

import os
from transagent.backend.config import get_config


def export_to_format(md_text: str, target_format: str,
                     assets_dir: str = "",
                     aligned_pairs=None,
                     placeholder_map=None) -> str:
    """
    按需导出。

    Args:
        md_text: 翻译后·已还原占位符的MD文本
        target_format: "md" | "docx" | "html" | "bilingual"
        assets_dir: assets目录路径（用于嵌入图片）
        aligned_pairs: 句级对齐结果（list[AlignedPair]·仅双语导出使用）
        placeholder_map: 占位符映射表（仅双语导出·用于还原源文中的占位符）

    Returns:
        导出文件路径
    """
    cfg = get_config().app
    output_dir = os.path.join(cfg.workspace_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    if target_format == "md":
        return _export_md(md_text, output_dir)
    elif target_format == "docx":
        return _export_docx(md_text, output_dir, assets_dir)
    elif target_format == "html":
        return _export_html(md_text, output_dir)
    elif target_format == "bilingual":
        return _export_bilingual(md_text, output_dir, aligned_pairs, placeholder_map)
    else:
        raise ValueError(f"不支持的导出格式: {target_format}")


def _export_md(md_text: str, output_dir: str) -> str:
    """MD→原样导出（UTF-8 纯文本）。

    整合方向为「MD中心制」：MD 是系统内唯一流通格式，原样导出零转换损失，
    便于二次编辑与版本管理。
    """
    output_path = os.path.join(output_dir, "output.md")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    return output_path


def _export_docx(md_text: str, output_dir: str, assets_dir: str) -> str:
    """MD→docx反向重建"""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    lines = md_text.split('\n')

    for line in lines:
        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level <= 6:
                doc.add_heading(text, level=level)
            else:
                doc.add_paragraph(text)

        # 图片引用
        elif line.startswith('![') and '](' in line:
            import re
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(5))
                    if alt_text:
                        doc.add_paragraph(f'[图] {alt_text}')

        # 代码块（简化处理·标注为代码）
        elif line.startswith('```'):
            continue  # 代码块由前后```标记，中间内容按普通段落处理
        elif line.strip().startswith('`') and line.strip().endswith('`'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('`'))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)

        # 空行
        elif not line.strip():
            continue

        # 普通段落
        else:
            doc.add_paragraph(line)

    output_path = os.path.join(output_dir, "output.docx")
    doc.save(output_path)
    return output_path


def _export_html(md_text: str, output_dir: str) -> str:
    """MD→HTML（简易渲染·生产环境可用marked.js等前端渲染）"""
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>翻译结果</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
               max-width: 800px; margin: 40px auto; padding: 20px;
               line-height: 1.8; color: #333; }}
        h1,h2,h3 {{ color: #1a1a1a; margin-top: 1.5em; }}
        code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px;
               font-family: "Fira Code", Consolas, monospace; }}
        pre {{ background: #2d2d2d; color: #f8f8f2; padding: 16px;
             border-radius: 8px; overflow-x: auto; }}
        pre code {{ background: transparent; padding: 0; color: inherit; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th,td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background: #f5f5f5; }}
        img {{ max-width: 100%; }}
    </style>
</head>
<body>
    <pre style="background:transparent;color:inherit;padding:0;white-space:pre-wrap;">{md_text}</pre>
</body>
</html>"""
    output_path = os.path.join(output_dir, "output.html")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
    return output_path


def _export_bilingual(md_text: str, output_dir: str,
                      aligned_pairs=None, placeholder_map=None) -> str:
    """
    双语对照导出（docx·源文/译文左右对照表格）。

    使用句级对齐结果 aligned_pairs 逐句对照：
      - 每个句对占一行：左列原文、右列译文
      - 标题块（#开头）合并为整行小节标题
      - 源文中的占位符用 placeholder_map 还原为可读文本
    无对齐句对时降级为纯译文docx。

    Args:
        md_text: 翻译后·已还原占位符的MD文本
        output_dir: 输出目录
        aligned_pairs: 句级对齐结果（list[AlignedPair]·可空·空则降级）
        placeholder_map: 占位符映射表（可空）

    Returns:
        docx文件路径
    """
    from docx import Document
    from docx.shared import Pt

    from transagent.backend.pipeline.restore import restore_placeholders

    def _has_content(pair) -> bool:
        # D8：双语对照要求源/译都有内容——合并块的"空译行"（源句译文已并入上一行）
        # 不单独成行，避免译文列出现空洞
        return (pair is not None
                and (getattr(pair, "source_seg", "") or "").strip()
                and (getattr(pair, "target_seg", "") or "").strip())

    def _restore(text: str) -> str:
        if placeholder_map and text:
            return restore_placeholders(text, placeholder_map)
        return text

    pairs = [p for p in (aligned_pairs or []) if _has_content(p)]

    if not pairs:
        # 降级：无对齐句对时导出纯译文docx
        return _export_docx(md_text, output_dir, "")

    doc = Document()
    doc.add_heading("双语对照翻译", level=0)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # 表头
    header = table.rows[0].cells
    header[0].text = "原文"
    header[1].text = "译文"
    for cell in header:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for pair in pairs:
        src = _restore(pair.source_seg or "")
        tgt = _restore(pair.target_seg or "")

        if src.lstrip().startswith("#"):
            # 标题块 → 合并整行作为小节标题
            row = table.add_row()
            merged = row.cells[0].merge(row.cells[1])
            merged.text = src.lstrip("#").strip()
            for para in merged.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(12)
        else:
            row = table.add_row()
            row.cells[0].text = src
            row.cells[1].text = tgt

    output_path = os.path.join(output_dir, "bilingual.docx")
    doc.save(output_path)
    return output_path
